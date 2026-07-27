# -*- coding: utf-8 -*-
"""md2word — Markdown 转 Word (.docx)，遵循个人排版习惯：

- `**文本**` 必须转为真正的加粗，不允许残留字面 **
- 不输出任何分隔符：跳过 --- / *** / ___ 水平线，不加分页符，
  表格的 |---| 分隔行丢弃（表格是原生 Word 表格），引用块转普通段落
- markdown 表格转原生 Word 表格（Table Grid 边框，表头加粗）
- 标题：# Title / ## Heading 1 / ### Heading 2 / #### Heading 3
- 列表：- 无序（按缩进嵌套）、1. 有序，使用原生列表样式
- 代码块与 `行内代码` 用 Consolas 等宽
- 中文正文微软雅黑，西文 Calibri

用法：python3 md2word.py <input.md> [output.docx]
依赖：python-docx（pip3 install python-docx）
"""
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    sys.stderr.write("缺少依赖 python-docx，请先执行: pip3 install python-docx\n")
    sys.exit(1)

EAST_FONT = u'微软雅黑'
MONO_FONT = 'Consolas'
TOKEN_RE = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
HR_RE = re.compile(r'^\s*(-{3,}|\*{3,}|_{3,})\s*$')


def style_run(run, mono=False):
    if mono:
        run.font.name = MONO_FONT
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), MONO_FONT if mono else EAST_FONT)


def add_runs(paragraph, text, force_bold=False):
    """解析行内格式：** 转加粗 run，` 转等宽 run。"""
    for part in TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            style_run(r)
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            style_run(r, mono=True)
        else:
            r = paragraph.add_run(part)
            if force_bold:
                r.bold = True
            style_run(r)


def parse_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def is_sep_row(line):
    return re.match(r'^\|[\s:\-|]+\|?\s*$', line.strip()) is not None


HEADING_STYLES = {1: 'Title', 2: 'Heading 1', 3: 'Heading 2', 4: 'Heading 3'}
BULLET_STYLES = {0: 'List Bullet', 1: 'List Bullet 2', 2: 'List Bullet 3'}


def convert(src, dst):
    with open(src, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), EAST_FONT)

    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            style_run(r, mono=True)
            r.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        if not stripped or HR_RE.match(stripped):
            i += 1
            continue

        # 表格：连续 | 开头的行，跳过分隔行，转原生 Word 表格
        if stripped.startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            rows = [parse_row(l) for l in tbl_lines if not is_sep_row(l)]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            add_runs(table.rows[ri].cells[ci].paragraphs[0],
                                     cell, force_bold=(ri == 0))
            continue

        m = re.match(r'^(#{1,4})\s+', stripped)
        if m:
            add_runs(doc.add_paragraph(style=HEADING_STYLES[len(m.group(1))]),
                     stripped[m.end():])
        elif stripped.startswith('>'):
            # 引用块转普通段落，不带任何引用样式（避免边框类装饰）
            add_runs(doc.add_paragraph(), stripped.lstrip('>').strip())
        elif re.match(r'^\d+\.\s+', stripped):
            add_runs(doc.add_paragraph(style='List Number'),
                     re.sub(r'^\d+\.\s+', '', stripped))
        elif re.match(r'^\s*[-*+]\s+', line):
            level = min((len(line) - len(line.lstrip())) // 2, 2)
            add_runs(doc.add_paragraph(style=BULLET_STYLES[level]),
                     re.sub(r'^\s*[-*+]\s+', '', line))
        else:
            add_runs(doc.add_paragraph(), stripped)
        i += 1

    doc.save(dst)
    return dst


def main():
    if len(sys.argv) < 2:
        sys.stderr.write('用法: python3 md2word.py <input.md> [output.docx]\n')
        sys.exit(1)
    src = sys.argv[1]
    if not os.path.isfile(src):
        sys.stderr.write('文件不存在: %s\n' % src)
        sys.exit(1)
    dst = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(src)[0] + '.docx'
    convert(src, dst)
    print('saved: %s' % dst)


if __name__ == '__main__':
    main()
